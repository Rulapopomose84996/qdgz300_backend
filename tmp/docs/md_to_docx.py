from __future__ import annotations

import re
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from latex2mathml.converter import convert as latex_to_mathml
from lxml import etree


ROOT = Path(r"D:\WorkSpace\Company\Tower\qdgz300_backend")
SOURCE_MD = ROOT / "docs" / "Inbox" / "待整理" / "研制任务书" / "信号处理软件研制任务书.md"
OUTPUT_DOCX = ROOT / "output" / "doc" / "信号处理软件研制任务书.docx"
XSL_PATH = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
INLINE_MATH_RE = re.compile(r"\$(.+?)\$")
BLOCK_MATH_RE = re.compile(r"\$\$\s*(.*?)\s*\$\$", re.S)

TRANSFORM = etree.XSLT(etree.parse(str(XSL_PATH)))


def latex_to_omml(latex: str):
    mathml = latex_to_mathml(latex)
    tree = etree.fromstring(mathml.encode("utf-8"))
    omml = TRANSFORM(tree)
    return etree.fromstring(etree.tostring(omml.getroot()))


def set_east_asia_font(run, east_asia_name: str):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia_name)


def set_run_style(run, *, bold=None, italic=None, code=False, link=False):
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if code:
        run.font.name = "Consolas"
        set_east_asia_font(run, "等线")
        run.font.size = Pt(10.5)
    else:
        run.font.name = "Times New Roman"
        set_east_asia_font(run, "宋体")
        run.font.size = Pt(12)
    if link:
        run.underline = True


def configure_paragraph(paragraph, *, first_line_indent=True, spacing_before=0, spacing_after=6, line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(spacing_before)
    fmt.space_after = Pt(spacing_after)
    fmt.line_spacing = line_spacing
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.first_line_indent = Cm(0.74) if first_line_indent else Cm(0)


def add_inline_text_with_math(paragraph, text: str, *, bold=False, italic=False):
    last = 0
    for match in INLINE_MATH_RE.finditer(text):
        prefix = text[last:match.start()]
        if prefix:
            run = paragraph.add_run(prefix)
            set_run_style(run, bold=bold, italic=italic)
        paragraph._element.append(latex_to_omml(match.group(1).strip()))
        last = match.end()
    suffix = text[last:]
    if suffix:
        run = paragraph.add_run(suffix)
        set_run_style(run, bold=bold, italic=italic)


def add_inline(paragraph, node, *, bold=False, italic=False):
    if isinstance(node, str):
        add_inline_text_with_math(paragraph, node, bold=bold, italic=italic)
        return
    if isinstance(node, NavigableString):
        add_inline_text_with_math(paragraph, str(node), bold=bold, italic=italic)
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    if name == "br":
        paragraph.add_run().add_break(WD_BREAK.LINE)
        return
    if name in {"strong", "b"}:
        for child in node.children:
            add_inline(paragraph, child, bold=True, italic=italic)
        return
    if name in {"em", "i"}:
        for child in node.children:
            add_inline(paragraph, child, bold=bold, italic=True)
        return
    if name == "code":
        run = paragraph.add_run(node.get_text())
        set_run_style(run, bold=bold, italic=italic, code=True)
        return
    if name == "a":
        run = paragraph.add_run(node.get_text())
        set_run_style(run, bold=bold, italic=italic, link=True)
        return
    for child in node.children:
        add_inline(paragraph, child, bold=bold, italic=italic)


def add_heading(doc: Document, text_value: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {min(level, 4)}"
    configure_paragraph(p, first_line_indent=False, spacing_before=12 if level <= 2 else 6, spacing_after=6, line_spacing=1.2)
    run = p.add_run(text_value)
    run.bold = True
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "黑体")
    run.font.size = {1: Pt(16), 2: Pt(14), 3: Pt(13), 4: Pt(12)}.get(level, Pt(12))


def add_paragraph_from_tag(doc: Document, tag: Tag):
    p = doc.add_paragraph()
    configure_paragraph(p, first_line_indent=True)
    for child in tag.children:
        add_inline(p, child)
    return p


def add_list(doc: Document, list_tag: Tag, level=0):
    ordered = list_tag.name.lower() == "ol"
    for li in list_tag.find_all("li", recursive=False):
        p = doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        configure_paragraph(p, first_line_indent=False, spacing_after=3, line_spacing=1.35)
        p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
        p.paragraph_format.first_line_indent = Cm(0)
        for child in li.contents:
            if isinstance(child, Tag) and child.name.lower() in {"ul", "ol"}:
                add_list(doc, child, level + 1)
            elif isinstance(child, Tag) and child.name.lower() == "p":
                for grand in child.children:
                    add_inline(p, grand)
            else:
                add_inline(p, child)


def fill_cell(cell, tag: Tag):
    cell.text = ""
    p = cell.paragraphs[0]
    configure_paragraph(p, first_line_indent=False, spacing_after=0, line_spacing=1.2)
    for child in tag.children:
        if isinstance(child, Tag) and child.name.lower() == "p":
            for grand in child.children:
                add_inline(p, grand)
        else:
            add_inline(p, child)
    for run in p.runs:
        if not run.font.size:
            run.font.size = Pt(10.5)
            run.font.name = "Times New Roman"
            set_east_asia_font(run, "宋体")


def add_table(doc: Document, table_tag: Tag):
    rows = table_tag.find_all("tr")
    matrix = []
    max_cols = 0
    for tr in rows:
        cells = tr.find_all(["th", "td"], recursive=False) or tr.find_all(["th", "td"])
        matrix.append(cells)
        max_cols = max(max_cols, len(cells))
    table = doc.add_table(rows=len(matrix), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, cells in enumerate(matrix):
        for c_idx, source in enumerate(cells):
            fill_cell(table.cell(r_idx, c_idx), source)
        if r_idx == 0:
            for cell in table.rows[r_idx].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph("")


def add_hr(doc: Document):
    p = doc.add_paragraph()
    configure_paragraph(p, first_line_indent=False, spacing_before=3, spacing_after=3, line_spacing=1.0)
    p_pr = p._element.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    border.append(bottom)
    p_pr.append(border)


def add_block_formula(doc: Document, latex: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    configure_paragraph(p, first_line_indent=False, spacing_before=3, spacing_after=3, line_spacing=1.0)
    p._element.append(latex_to_omml(latex.strip()))


def process_node(doc: Document, node, first_h1_processed=False):
    if isinstance(node, NavigableString):
        return first_h1_processed
    if not isinstance(node, Tag):
        return first_h1_processed
    name = node.name.lower()
    if name == "div" and "math-block" in (node.get("class") or []):
        add_block_formula(doc, node.get_text())
        return first_h1_processed
    if name == "h1":
        if not first_h1_processed:
            return True
        add_heading(doc, node.get_text(strip=True), 1)
        return first_h1_processed
    if name in {"h2", "h3", "h4", "h5", "h6"}:
        add_heading(doc, node.get_text(strip=True), int(name[1]))
        return first_h1_processed
    if name == "p":
        add_paragraph_from_tag(doc, node)
        return first_h1_processed
    if name in {"ul", "ol"}:
        add_list(doc, node)
        return first_h1_processed
    if name == "table":
        add_table(doc, node)
        return first_h1_processed
    if name == "hr":
        add_hr(doc)
        return first_h1_processed
    for child in node.children:
        first_h1_processed = process_node(doc, child, first_h1_processed)
    return first_h1_processed


def protect_block_math(text: str) -> str:
    def repl(match):
        formula = match.group(1).strip()
        return f'\n<div class="math-block">{formula}</div>\n'

    return BLOCK_MATH_RE.sub(repl, text)


def build_doc(source_text: str):
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    html = markdown.markdown(protect_block_math(source_text), extensions=["tables", "fenced_code", "sane_lists"])
    soup = BeautifulSoup(f"<div>{html}</div>", "lxml")
    root_tag = soup.div

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for style_name, east_asia, size in [
        ("Heading 1", "黑体", Pt(16)),
        ("Heading 2", "黑体", Pt(14)),
        ("Heading 3", "黑体", Pt(13)),
        ("Heading 4", "黑体", Pt(12)),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = size
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)

    title_match = re.search(r"^#\s+(.+)$", source_text, re.M)
    title = title_match.group(1).strip() if title_match else SOURCE_MD.stem
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    set_east_asia_font(title_run, "黑体")
    title_run.font.size = Pt(20)
    configure_paragraph(title_p, first_line_indent=False, spacing_before=0, spacing_after=12, line_spacing=1.0)
    doc.core_properties.title = title

    first_h1 = False
    for child in root_tag.children:
        first_h1 = process_node(doc, child, first_h1)

    doc.save(OUTPUT_DOCX)
    print(f"output={OUTPUT_DOCX}")


if __name__ == "__main__":
    build_doc(SOURCE_MD.read_text(encoding="utf-8"))
